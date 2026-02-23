import streamlit as st
import os
from groq import Groq
import google.generativeai as genai
from docx import Document
import io

# Nastavení vzhledu
st.set_page_config(page_title="Chytrý zápis ze schůzky", page_icon="📝", layout="centered")

# Načtení klíčů z tajného trezoru
try:
    groq_api_key = st.secrets["GROQ_API_KEY"]
    gemini_api_key = st.secrets["GEMINI_API_KEY"]
except:
    st.error("Chybí API klíče v nastavení aplikace.")
    st.stop()

# --- PAMĚŤ APLIKACE ---
if "transcription" not in st.session_state:
    st.session_state.transcription = None
if "zapis_text" not in st.session_state:
    st.session_state.zapis_text = None
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

st.title("📝 Generátor manažerských zápisů")
st.markdown("Nahraj audio ze schůzky a vyber si, jakou šablonu zápisu chceš vygenerovat. Dole se pak můžeš doptávat v chatu.")

st.info("⚠️ **Limit velikosti souboru:** Maximálně **25 MB**. Větší soubory zmenši zdarma zde: [Compress audio online](https://www.freeconvert.com/audio-compressor)")

# Nahrání souboru
audio_file = st.file_uploader("Nahraj záznam ze schůzky (MP3, WAV, M4A)", type=['mp3', 'wav', 'm4a'])

# --- DVĚ TLAČÍTKA VEDLE SEBE ---
col_btn1, col_btn2 = st.columns(2)
with col_btn1:
    btn_standard = st.button("🚀 Vygenerovat standardní zápis", use_container_width=True)
with col_btn2:
    btn_obecny = st.button("📋 Vygenerovat obecný zápis", use_container_width=True)

# Spustí se, pokud uživatel klikne na JAKÉKOLIV z tlačítek
if btn_standard or btn_obecny:
    if not audio_file:
        st.warning("Nejprve prosím nahraj soubor s audiem.")
    else:
        try:
            st.session_state.chat_history = []
            
            # 1. PŘEPIS AUDIA
            with st.spinner("⏳ Poslouchám a přepisuji audio (může to chvilku trvat)..."):
                with open("temp_audio.mp3", "wb") as f:
                    f.write(audio_file.getbuffer())
                
                client = Groq(api_key=groq_api_key)
                with open("temp_audio.mp3", "rb") as file:
                    vysledek_prepisu = client.audio.transcriptions.create(
                      file=("temp_audio.mp3", file.read()),
                      model="whisper-large-v3",
                      response_format="text",
                      language="cs"
                    )
                os.remove("temp_audio.mp3")
                st.session_state.transcription = vysledek_prepisu
            
            st.success("✅ Přepis byl úspěšně dokončen!")

            # 2. TVORBA ZÁPISU PODLE VYBRANÉ ŠABLONY
            with st.spinner("⏳ Generuji zápis podle vybrané šablony..."):
                genai.configure(api_key=gemini_api_key)
                model = genai.GenerativeModel('gemini-2.5-flash')
                
                # Pokud klikl na první tlačítko
                if btn_standard:
                    prompt = f"""
                    Jsi profesionální firemní asistent. Přečti si následující surový přepis ze schůzky a vytvoř z něj přehledný manažerský zápis v češtině.
                    Rozděl ho na:
                    1. Hlavní téma schůzky
                    2. Nejdůležitější probrané body (odrážky)
                    3. Učiněná rozhodnutí
                    4. Akční kroky / Úkoly (Kdo má co udělat)
                    
                    Zde je přepis:
                    {st.session_state.transcription}
                    """
                
                # Pokud klikl na druhé tlačítko ("Obecný zápis")
                elif btn_obecny:
                    prompt = f"""
                    Jsi profesionální firemní asistent. Přečti si následující surový přepis ze schůzky a vytvoř z něj přesný zápis v češtině PŘESNĚ podle následující šablony. 
                    Dodržuj formátování (nadpisy, tučné písmo) a řiď se instrukcemi, které jsou uvedeny v hranatých závorkách.

                    ## MANAZERSKE SHRNUTÍ
                    **Cíl setkání:** [jedna az dve vety]
                    **Klícová rozhodnutí:** [kazde rozhodnutí na novy radek s pomlckou;pokud zadne nepadlo napís: Bez formalnich rozhodnutí]
                    ---
                    ## DISKUTOVANÁ TÉMATA
                    [kazde tema na novy radek s pomlckou, max 8 bodu]
                    ---
                    ## AKCNÍ KROKY
                    | # | Úkol | Odpovědná osoba | Termín | Stav |
                    |---|------|-----------------|--------|------|
                    [radky tabulky; pokud neni termin nebo osoba napís Neurčeno; Stav vzdy Nový]

                    Zde je přepis:
                    {st.session_state.transcription}
                    """

                response = model.generate_content(prompt)
                st.session_state.zapis_text = response.text
                
        except Exception as e:
            st.error(f"Ouvej, něco se pokazilo: {e}")

# --- ZOBRAZENÍ VÝSLEDKŮ A CHATU ---
if st.session_state.transcription and st.session_state.zapis_text:
    
    st.success("✅ Zápis je hotový!")
    st.markdown("### Náhled zápisu:")
    st.write(st.session_state.zapis_text)

    # 3. TVORBA WORD DOKUMENTŮ PRO STAŽENÍ
    st.markdown("### 💾 Ke stažení:")
    col1, col2 = st.columns(2)
    
    with col1:
        doc_zapis = Document()
        doc_zapis.add_heading('Zápis ze schůzky', 0)
        doc_zapis.add_paragraph(st.session_state.zapis_text)
        bio_zapis = io.BytesIO()
        doc_zapis.save(bio_zapis)
        
        st.download_button(
            label="📝 Stáhnout zápis",
            data=bio_zapis.getvalue(),
            file_name="zapis_ze_schuzky.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        
    with col2:
        doc_prepis = Document()
        doc_prepis.add_heading('Kompletní přepis schůzky', 0)
        doc_prepis.add_paragraph(st.session_state.transcription)
        bio_prepis = io.BytesIO()
        doc_prepis.save(bio_prepis)
        
        st.download_button(
            label="🗣️ Stáhnout doslovný přepis",
            data=bio_prepis.getvalue(),
            file_name="kompletni_prepis.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    # 4. CHATOVÁNÍ S PŘEPISEM
    st.markdown("---")
    st.markdown("### 💬 Zeptejte se na detaily ze schůzky")
    st.caption("Chybí vám v zápisu něco? Napište otázku a umělá inteligence to v textu dohledá.")

    for message in st.session_state.chat_history:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    if user_question := st.chat_input("Zeptejte se... (např. 'Jaký byl dohodnutý termín spuštění?')"):
        st.session_state.chat_history.append({"role": "user", "content": user_question})
        with st.chat_message("user"):
            st.markdown(user_question)

        with st.chat_message("assistant"):
            with st.spinner("Dohledávám v přepisu..."):
                genai.configure(api_key=gemini_api_key)
                model = genai.GenerativeModel('gemini-2.5-flash')
                
                chat_prompt = f"""
                Jsi asistent. Tvojí jedinou prací je odpovídat na otázky týkající se této schůzky, POUZE na základě poskytnutého přepisu.
                Pokud odpověď v přepisu nenajdeš, omluv se a řekni: "Tato informace v přepisu bohužel nezazněla."
                
                Zde je přepis schůzky:
                {st.session_state.transcription}
                
                Zde je otázka uživatele:
                {user_question}
                """
                
                odpoved = model.generate_content(chat_prompt)
                st.markdown(odpoved.text)
                st.session_state.chat_history.append({"role": "assistant", "content": odpoved.text})
